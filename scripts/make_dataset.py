"""Generate the Milestone 2 dataset: 30+ diverse resumes and 6 job descriptions.

Extends Milestone 1's ``make_samples.py`` with a richer, *labelled* dataset for
the RAG pipeline:

* 36 resumes across 11 role families, in a mix of .txt / .pdf / .docx
* three section-header styles and two section orders, so the section-aware
  chunker in ``resume_rag.py`` is exercised against real variety
* every resume has SKILLS, dated EXPERIENCE and EDUCATION sections, so
  metadata extraction can be evaluated
* ``dataset/labels.json`` records ground truth (role family, years of
  experience, skills, education level) for every file, plus which candidate
  roles are relevant to each job description — the analysis notebook uses this
  to compute retrieval accuracy

The 8 Milestone 1 resumes keep their filenames (john_doe.pdf, ...) so the old
README examples still work; their content is upgraded to the richer format.

Deterministic: seeded RNG and a fixed "today" year. Run from the project root:

    python scripts/make_dataset.py

Hard mode (messier corpus for stress-testing extractors):

    python scripts/make_dataset.py --hard
"""

from __future__ import annotations

import argparse
import json
import random
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

PROJECT_ROOT = Path(__file__).resolve().parent.parent
RESUMES_DIR = PROJECT_ROOT / "resumes"
JOBS_DIR = PROJECT_ROOT / "job_descriptions"
DATASET_DIR = PROJECT_ROOT / "dataset"

SEED = 42
TODAY_YEAR = 2026  # fixed so generated dates and labels never drift

# --- Hard mode ---
HARD_SEED = 4242
RESUMES_HARD_DIR = PROJECT_ROOT / "resumes_hard"
TODAY_MONTH = 6  # mid-year anchor, matches extract_metadata's today_year pin

# --- Role catalogue -----------------------------------------------------------

ROLES: Dict[str, Dict[str, object]] = {
    "ml": {
        "titles": ["Machine Learning Engineer", "Data Scientist", "Applied Scientist"],
        "core": ["Python", "PyTorch", "scikit-learn", "pandas", "SQL", "Machine Learning"],
        "extra": ["TensorFlow", "NLP", "Computer Vision", "MLOps", "Deep Learning", "Airflow", "Docker", "AWS", "Spark"],
        "field": ["Computer Science", "Data Science", "Statistics"],
        "bullets": [
            "Built and deployed {s} models improving prediction accuracy by {pct}%.",
            "Productionised {s} pipelines serving {n}M inferences per day.",
            "Led feature engineering with {s} and {s2} across {n} data sources.",
            "Reduced model training time {pct}% by optimising {s} workflows.",
            "Shipped an A/B-tested recommendation system using {s} and {s2}.",
        ],
        "projects": [
            "Churn Predictor: end-to-end {s} pipeline with automated retraining.",
            "Resume Screener: NLP classifier built with {s} and {s2}.",
        ],
        "certs": ["AWS Certified Machine Learning - Specialty", "TensorFlow Developer Certificate"],
    },
    "backend_python": {
        "titles": ["Backend Engineer", "Python Developer", "Software Engineer, Backend"],
        "core": ["Python", "FastAPI", "PostgreSQL", "Docker", "REST APIs"],
        "extra": ["Django", "Flask", "Redis", "Kafka", "AWS", "Microservices", "Kubernetes", "CI/CD", "MongoDB"],
        "field": ["Computer Science", "Software Engineering"],
        "bullets": [
            "Designed {s} microservices handling {n}M requests per day.",
            "Cut p95 API latency {pct}% by tuning {s} queries and {s2} caching.",
            "Built event-driven workflows with {s} and {s2}.",
            "Owned CI/CD pipelines deploying {s} services {n}x per week.",
            "Migrated a monolith to {s} microservices with zero-downtime releases.",
        ],
        "projects": [
            "Billing API: idempotent payment service in {s} backed by {s2}.",
            "Rate Limiter: distributed limiter using {s} and {s2}.",
        ],
        "certs": ["AWS Certified Developer - Associate"],
    },
    "backend_java": {
        "titles": ["Java Backend Developer", "Software Engineer, Java", "Senior Java Engineer"],
        "core": ["Java", "Spring Boot", "MySQL", "Kafka", "REST APIs"],
        "extra": ["Microservices", "Docker", "Kubernetes", "AWS", "Redis", "PostgreSQL", "Maven"],
        "field": ["Computer Science", "Information Systems"],
        "bullets": [
            "Designed {s} services processing {n}M transactions daily.",
            "Improved throughput {pct}% by introducing {s} message streaming.",
            "Hardened {s} APIs with circuit breakers and structured retries.",
            "Modernised legacy modules into {s} microservices on {s2}.",
        ],
        "projects": [
            "Ledger Service: double-entry bookkeeping engine on {s} and {s2}.",
        ],
        "certs": ["Oracle Certified Professional: Java SE"],
    },
    "frontend": {
        "titles": ["Frontend Engineer", "UI Engineer", "Frontend Developer"],
        "core": ["JavaScript", "TypeScript", "React", "CSS", "HTML"],
        "extra": ["Next.js", "Tailwind CSS", "GraphQL", "Vue.js", "Cypress", "Redux"],
        "field": ["Computer Science", "Software Engineering"],
        "bullets": [
            "Built a {s} design system used across {n} product teams.",
            "Improved Core Web Vitals {pct}% by optimising {s} rendering.",
            "Shipped accessible {s} components meeting WCAG 2.1 AA.",
            "Migrated a legacy UI to {s} with {s2}, cutting bundle size {pct}%.",
        ],
        "projects": [
            "Dashboard Kit: reusable {s} charting components with {s2}.",
        ],
        "certs": [],
    },
    "fullstack": {
        "titles": ["Full-Stack Engineer", "Full Stack Developer", "Product Engineer"],
        "core": ["Python", "React", "TypeScript", "PostgreSQL", "Docker"],
        "extra": ["Django", "FastAPI", "Node.js", "Redis", "AWS", "GraphQL", "Next.js"],
        "field": ["Computer Science", "Software Engineering"],
        "bullets": [
            "Delivered {s} + {s2} features end to end for {n}0k users.",
            "Built REST and GraphQL APIs in {s} consumed by {s2} clients.",
            "Reduced page load {pct}% with {s} caching and query tuning.",
            "Owned the {s} stack from schema design to {s2} UI.",
        ],
        "projects": [
            "Marketplace MVP: {s} backend with a {s2} storefront.",
        ],
        "certs": [],
    },
    "devops": {
        "titles": ["DevOps Engineer", "Platform Engineer", "Site Reliability Engineer"],
        "core": ["Kubernetes", "Terraform", "AWS", "Docker", "CI/CD", "Linux"],
        "extra": ["Python", "Bash", "Prometheus", "Grafana", "Ansible", "GCP", "Helm"],
        "field": ["Computer Science", "Information Systems"],
        "bullets": [
            "Automated {s} infrastructure with {s2}, cutting provisioning time {pct}%.",
            "Ran {n}0+ {s} clusters with GitOps-driven deployments.",
            "Cut deployment lead time {pct}% by rebuilding {s} pipelines.",
            "Built {s} observability dashboards reducing MTTR {pct}%.",
        ],
        "projects": [
            "Self-Service Platform: {s} modules enabling one-click environments.",
        ],
        "certs": ["Certified Kubernetes Administrator (CKA)", "AWS Certified Solutions Architect - Associate"],
    },
    "data_engineer": {
        "titles": ["Data Engineer", "Analytics Engineer", "Big Data Engineer"],
        "core": ["Python", "SQL", "Spark", "Airflow", "ETL"],
        "extra": ["Snowflake", "dbt", "Kafka", "AWS", "PostgreSQL", "Data Warehousing"],
        "field": ["Computer Science", "Data Science", "Information Systems"],
        "bullets": [
            "Built {s} pipelines moving {n}TB per day into the warehouse.",
            "Cut pipeline failures {pct}% with {s} data-quality checks.",
            "Modelled analytics tables in {s} powering {n}0+ dashboards.",
            "Migrated batch ETL to streaming with {s} and {s2}.",
        ],
        "projects": [
            "CDC Pipeline: change-data-capture into {s} using {s2}.",
        ],
        "certs": ["Google Cloud Professional Data Engineer"],
    },
    "mobile": {
        "titles": ["Mobile Engineer", "Android Developer", "iOS Developer"],
        "core": ["Kotlin", "Android", "REST APIs", "Git"],
        "extra": ["Swift", "iOS", "React Native", "Flutter", "Firebase", "GraphQL"],
        "field": ["Computer Science", "Software Engineering"],
        "bullets": [
            "Shipped {s} apps with {n}00k+ installs and 4.7-star ratings.",
            "Cut app startup time {pct}% by profiling {s} initialisation.",
            "Built offline-first sync with {s} and {s2}.",
            "Led migration from {s2} to {s}, reducing crash rate {pct}%.",
        ],
        "projects": [
            "Fitness Tracker: {s} app with {s2} cloud sync.",
        ],
        "certs": [],
    },
    "qa": {
        "titles": ["QA Automation Engineer", "Software Engineer in Test"],
        "core": ["Selenium", "Python", "pytest", "CI/CD", "REST APIs"],
        "extra": ["Cypress", "Playwright", "Java", "JUnit", "Docker"],
        "field": ["Computer Science", "Information Systems"],
        "bullets": [
            "Built a {s} regression suite covering {n}00+ scenarios.",
            "Cut flaky tests {pct}% by stabilising {s} fixtures.",
            "Automated API contract tests with {s} and {s2}.",
        ],
        "projects": [
            "Test Harness: parallel {s} runner integrated into {s2}.",
        ],
        "certs": ["ISTQB Certified Tester"],
    },
    "security": {
        "titles": ["Security Engineer", "Application Security Engineer"],
        "core": ["Penetration Testing", "OWASP", "Burp Suite", "Linux", "Python"],
        "extra": ["SIEM", "AWS", "Docker", "Threat Modeling"],
        "field": ["Cybersecurity", "Computer Science"],
        "bullets": [
            "Ran {n}0+ penetration tests; remediated all critical findings.",
            "Cut mean time to detect {pct}% by tuning {s} alerting.",
            "Embedded {s} checks into CI, blocking vulnerable releases.",
        ],
        "projects": [
            "AppSec Pipeline: automated {s} scanning with {s2} triage.",
        ],
        "certs": ["OSCP", "CISSP"],
    },
    "design": {
        "titles": ["Product Designer", "UX Designer"],
        "core": ["Figma", "User Research", "Prototyping", "Design Systems"],
        "extra": ["HTML", "CSS", "Accessibility"],
        "field": ["Interaction Design", "Human-Computer Interaction"],
        "bullets": [
            "Ran {n}0+ usability studies shaping the product roadmap.",
            "Built a {s} design system adopted by {n} squads.",
            "Improved onboarding conversion {pct}% through {s} prototypes.",
        ],
        "projects": [
            "Design Tokens: cross-platform theming built in {s}.",
        ],
        "certs": [],
    },
}

# Candidates: (filename, display name, role, total years of experience).
# The first 8 keep their Milestone 1 filenames and approximate seniority.
CANDIDATES: List[Tuple[str, str, str, int]] = [
    ("john_doe.pdf", "John Doe", "backend_python", 8),
    ("jane_smith.docx", "Jane Smith", "frontend", 9),
    ("alex_lee.txt", "Alex Lee", "ml", 5),
    ("maria_garcia.pdf", "Maria Garcia", "devops", 9),
    ("sam_patel.txt", "Sam Patel", "backend_java", 10),
    ("lin_wei.docx", "Lin Wei", "ml", 6),
    ("olivia_brown.txt", "Olivia Brown", "design", 9),
    ("noah_kim.pdf", "Noah Kim", "fullstack", 6),
    ("priya_sharma.txt", "Priya Sharma", "ml", 7),
    ("daniel_okafor.pdf", "Daniel Okafor", "ml", 3),
    ("grace_chen.docx", "Grace Chen", "ml", 11),
    ("ethan_walker.txt", "Ethan Walker", "backend_python", 4),
    ("sofia_rossi.pdf", "Sofia Rossi", "backend_python", 6),
    ("omar_hassan.docx", "Omar Hassan", "backend_python", 12),
    ("ava_thompson.txt", "Ava Thompson", "backend_java", 5),
    ("lucas_martin.pdf", "Lucas Martin", "backend_java", 7),
    ("isabella_silva.docx", "Isabella Silva", "frontend", 4),
    ("ryan_oconnor.txt", "Ryan OConnor", "frontend", 6),
    ("mei_tanaka.pdf", "Mei Tanaka", "frontend", 2),
    ("arjun_mehta.docx", "Arjun Mehta", "fullstack", 8),
    ("hannah_weber.txt", "Hannah Weber", "fullstack", 3),
    ("david_cohen.pdf", "David Cohen", "fullstack", 11),
    ("fatima_khan.docx", "Fatima Khan", "devops", 6),
    ("carlos_mendez.txt", "Carlos Mendez", "devops", 4),
    ("emily_davis.pdf", "Emily Davis", "devops", 12),
    ("viktor_petrov.docx", "Viktor Petrov", "data_engineer", 8),
    ("nina_johansson.txt", "Nina Johansson", "data_engineer", 5),
    ("tomas_novak.pdf", "Tomas Novak", "data_engineer", 3),
    ("leila_aziz.docx", "Leila Aziz", "data_engineer", 10),
    ("jack_murphy.txt", "Jack Murphy", "mobile", 7),
    ("chloe_dubois.pdf", "Chloe Dubois", "mobile", 4),
    ("andre_williams.docx", "Andre Williams", "mobile", 9),
    ("yuki_sato.txt", "Yuki Sato", "qa", 6),
    ("elena_petrova.pdf", "Elena Petrova", "qa", 9),
    ("marcus_reed.docx", "Marcus Reed", "security", 8),
    ("aisha_bello.txt", "Aisha Bello", "security", 5),
]

COMPANIES = [
    "Acme Cloud", "DataWorks", "Pixel Labs", "Bright Apps", "Insight AI",
    "Quant Co", "Streamline", "HostPro", "FinServe", "RetailX", "Vision Corp",
    "Loop", "WebStart", "Craft Studio", "AppWorks", "Northwind Analytics",
    "BlueRiver Systems", "Helio Health", "Vertex Logistics", "Nimbus Pay",
    "Orchid Media", "Summit Robotics", "Clearpath Energy", "Mosaic Travel",
]

SCHOOLS = [
    "State University", "Tech Institute of Engineering", "Riverside University",
    "National Institute of Technology", "Northfield College", "Westlake University",
    "Global Tech University", "Lakeside State University",
]

CITIES = [
    "Austin, TX", "Seattle, WA", "Bengaluru, India", "Berlin, Germany",
    "Toronto, Canada", "London, UK", "Singapore", "Remote (US)", "Pune, India",
    "Amsterdam, Netherlands",
]

DEGREE_TEMPLATES = {
    "bachelor": ["B.S. in {field}", "Bachelor of Technology in {field}", "B.Sc. in {field}"],
    "master": ["M.S. in {field}", "Master of Science in {field}", "M.Tech in {field}"],
    "phd": ["Ph.D. in {field}"],
}

SECTION_STYLES = {
    "caps": {
        "summary": "SUMMARY", "skills": "SKILLS", "experience": "EXPERIENCE",
        "education": "EDUCATION", "projects": "PROJECTS", "certifications": "CERTIFICATIONS",
    },
    "title": {
        "summary": "Professional Summary", "skills": "Technical Skills",
        "experience": "Work Experience", "education": "Education",
        "projects": "Projects", "certifications": "Certifications",
    },
    "alt": {
        "summary": "PROFILE", "skills": "CORE COMPETENCIES",
        "experience": "EMPLOYMENT HISTORY", "education": "EDUCATION & TRAINING",
        "projects": "KEY PROJECTS", "certifications": "LICENSES & CERTIFICATIONS",
    },
}

SECTION_ORDERS = [
    ["summary", "skills", "experience", "projects", "education", "certifications"],
    ["summary", "skills", "education", "experience", "projects", "certifications"],
]


# --- Hard mode ----------------------------------------------------------------

# 40 new candidate names — ASCII, diverse, none reused from clean CANDIDATES.
# Tuple: (filename, display_name, role, total_years)
# Tier-2 slots are indices 28-39; we ensure 8 of those 12 are .txt so the
# no-skills-section check (raw text search) reliably finds >= 8 matches.
HARD_CANDIDATES: List[Tuple[str, str, str, int]] = [
    # indices 0-13  → tier 0
    ("rafael_reyes.txt",       "Rafael Reyes",       "ml",             6),
    ("amara_diallo.pdf",       "Amara Diallo",       "ml",             9),
    ("sven_lindqvist.docx",    "Sven Lindqvist",     "ml",             4),
    ("pooja_iyer.txt",         "Pooja Iyer",         "ml",             12),
    ("brendan_walsh.pdf",      "Brendan Walsh",      "backend_python", 5),
    ("yara_osei.docx",         "Yara Osei",          "backend_python", 8),
    ("nikolaj_berg.txt",       "Nikolaj Berg",       "backend_python", 3),
    ("lena_fischer.pdf",       "Lena Fischer",       "backend_python", 11),
    ("tatsuya_honda.docx",     "Tatsuya Honda",      "backend_java",   7),
    ("chisom_eze.txt",         "Chisom Eze",         "backend_java",   4),
    ("pawel_kowalski.pdf",     "Pawel Kowalski",     "backend_java",   10),
    ("ingrid_haugen.txt",      "Ingrid Haugen",      "backend_java",   2),
    ("mateus_carvalho.txt",    "Mateus Carvalho",    "frontend",       6),
    ("soraya_mansouri.pdf",    "Soraya Mansouri",    "frontend",       3),
    # indices 14-27 → tier 1
    ("kwame_asante.docx",      "Kwame Asante",       "frontend",       9),
    ("hana_novotna.txt",       "Hana Novotna",       "frontend",       5),
    ("dmitri_volkov.pdf",      "Dmitri Volkov",      "fullstack",      8),
    ("amelia_hunt.docx",       "Amelia Hunt",        "fullstack",      4),
    ("tariq_saleem.txt",       "Tariq Saleem",       "fullstack",      11),
    ("nora_eriksen.pdf",       "Nora Eriksen",       "fullstack",      3),
    ("jun_park.docx",          "Jun Park",           "devops",         7),
    ("fatou_fall.txt",         "Fatou Fall",         "devops",         5),
    ("anders_christensen.pdf", "Anders Christensen", "devops",         13),
    ("priti_nair.txt",         "Priti Nair",         "data_engineer",  6),
    ("tobias_braun.txt",       "Tobias Braun",       "data_engineer",  9),
    ("adaeze_obi.pdf",         "Adaeze Obi",         "data_engineer",  4),
    ("enrique_vega.docx",      "Enrique Vega",       "data_engineer",  14),
    ("haruto_yamamoto.txt",    "Haruto Yamamoto",    "mobile",         5),
    # indices 28-39 → tier 2 (8 .txt so no-skills check hits >=8)
    ("zoe_lambert.txt",        "Zoe Lambert",        "mobile",         8),
    ("olumide_adeyemi.txt",    "Olumide Adeyemi",    "mobile",         3),
    ("svetlana_ivanova.txt",   "Svetlana Ivanova",   "mobile",         10),
    ("cormac_byrne.txt",       "Cormac Byrne",       "qa",             6),
    ("meifen_zhong.pdf",       "Meifen Zhong",       "qa",             9),
    ("ibrahim_traore.txt",     "Ibrahim Traore",     "security",       7),
    ("petra_blum.txt",         "Petra Blum",         "security",       4),
    ("yusuf_hassan.txt",       "Yusuf Hassan",       "backend_python", 7),
    ("clara_stone.txt",        "Clara Stone",        "frontend",       2),
    ("rajan_pillai.pdf",       "Rajan Pillai",       "fullstack",      6),
    ("veronika_horvath.docx",  "Veronika Horvath",   "ml",             8),
    ("kweku_mensah.txt",       "Kweku Mensah",       "data_engineer",  5),
]

# Messy section-header mappings for tier-1 resumes.
MESSY_HEADERS: Dict[str, str] = {
    "experience": "Career History",
    "experience_alt": "Where I've Worked",
    "summary": "What I Bring",
    "skills": "Toolbox",
    "education": "Studies",
    "summary_alt": "My Background",
}

# Date style choices
DATE_STYLES = ["year", "monthname", "slash"]

# Present token variants cycled deterministically per resume index.
PRESENT_TOKENS = ["Present", "present", "Current"]

_MONTH_ABBR = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
               "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]


# --- Hard mode resume model ---------------------------------------------------


@dataclass(frozen=True)
class HardJob:
    title: str
    company: str
    start_year: int
    start_month: int    # 1-12
    end_year: int       # TODAY_YEAR for present
    end_month: int      # TODAY_MONTH for present
    is_present: bool
    bullets: Tuple[str, ...]


@dataclass(frozen=True)
class HardResume:
    filename: str
    name: str
    role: str
    title: str
    email: str
    phone: str
    location: str
    summary: str
    skills: Tuple[str, ...]
    jobs: Tuple[HardJob, ...]
    education_level: str    # "bachelor"/"master"/"phd"/"unknown"
    education_line: str     # "" when absent
    projects: Tuple[str, ...]
    certifications: Tuple[str, ...]
    style: str
    order: Tuple[str, ...]
    total_years: int        # floor(total career months / 12)
    tier: int               # 0/1/2
    date_style: str         # year/monthname/slash
    present_token: str      # Present/present/Current
    has_skills_section: bool
    has_summary_header: bool
    experience_header: str


# --- Hard mode helpers --------------------------------------------------------


def _make_hard_jobs(
    rng: random.Random,
    role: Dict[str, object],
    skills: Sequence[str],
    years: int,
) -> Tuple[HardJob, ...]:
    """Build contiguous job spans with explicit month granularity."""
    n_jobs = max(1, min(4, (years + 2) // 3))
    spans = [1] * n_jobs
    for _ in range(years - n_jobs):
        spans[rng.randrange(n_jobs)] += 1

    titles = list(role["titles"])  # type: ignore[index]
    bullets_pool = list(role["bullets"])  # type: ignore[index]
    companies = rng.sample(COMPANIES, n_jobs)

    career_start_month = rng.randint(1, 12)
    cur_year = TODAY_YEAR - years
    cur_month = career_start_month

    jobs: List[HardJob] = []
    for i, span in enumerate(spans):
        is_last = (i == n_jobs - 1)
        s_year, s_month = cur_year, cur_month
        if is_last:
            e_year, e_month = TODAY_YEAR, TODAY_MONTH
        else:
            e_year = cur_year + span
            e_month = cur_month
        title = rng.choice(titles)
        if is_last:
            title = _with_seniority(title, _seniority(years))
        job_bullets = tuple(
            _fill(b, rng, skills)
            for b in rng.sample(bullets_pool, min(2, len(bullets_pool)))
        )
        jobs.append(HardJob(
            title=title,
            company=companies[i],
            start_year=s_year,
            start_month=s_month,
            end_year=e_year,
            end_month=e_month,
            is_present=is_last,
            bullets=job_bullets,
        ))
        cur_year = e_year
        cur_month = e_month
    return tuple(reversed(jobs))


def _total_months(jobs: Tuple[HardJob, ...]) -> int:
    """Sum career span in months (start of earliest job → today)."""
    if not jobs:
        return 0
    earliest = min(jobs, key=lambda j: (j.start_year, j.start_month))
    total = (TODAY_YEAR - earliest.start_year) * 12 + (TODAY_MONTH - earliest.start_month)
    return max(total, 0)


def _skill_years_map(
    jobs: Tuple[HardJob, ...],
    skills: Tuple[str, ...],
    career_months: int,
    date_style: str,
) -> Dict[str, float]:
    """Compute skill_years for each skill.

    Merge intervals of jobs whose rendered text mentions the skill.
    Skills appearing in no job get the full career span (they live in
    skills/summary — the extractor mirrors this convention).
    """
    result: Dict[str, float] = {}
    for skill in skills:
        skill_lower = skill.lower()
        matched_months = 0
        for job in jobs:
            text = (job.title + " " + " ".join(job.bullets)).lower()
            if skill_lower in text:
                if date_style == "year":
                    span = (job.end_year - job.start_year) * 12
                else:
                    span = (job.end_year - job.start_year) * 12 + (job.end_month - job.start_month)
                matched_months += max(span, 0)
        if matched_months == 0:
            # skill lives in summary/skills line — assign full career span
            matched_months = career_months
        result[skill] = round(matched_months / 12, 1)
    return result


def _fmt_date(year: int, month: int, date_style: str, present_token: str, is_present: bool) -> str:
    if is_present:
        return present_token
    if date_style == "year":
        return str(year)
    if date_style == "monthname":
        return f"{_MONTH_ABBR[month - 1]} {year}"
    return f"{month:02d}/{year}"


def _rng_choice_from_list(choices: List[str], key: str) -> str:
    """Deterministic choice based on a string key (no rng needed)."""
    return choices[hash(key) % len(choices)]


def _make_hard_resume(idx: int, spec: Tuple[str, str, str, int], rng: random.Random) -> HardResume:
    filename, name, role_key, years = spec
    role = ROLES[role_key]
    core = list(role["core"])  # type: ignore[index]
    extra = list(role["extra"])  # type: ignore[index]
    skills = tuple(core + rng.sample(extra, min(len(extra), rng.randint(2, 4))))

    title_base = rng.choice(list(role["titles"]))  # type: ignore[index]
    display_title = _with_seniority(title_base, _seniority(years))

    # Tier assignment: target ~14 / 14 / 12
    tier_pool = [0] * 14 + [1] * 14 + [2] * 12
    tier = tier_pool[idx % len(tier_pool)]

    date_style = rng.choice(DATE_STYLES)
    present_token = PRESENT_TOKENS[idx % len(PRESENT_TOKENS)]

    # Education: tier-2 ~half absent
    has_education = not (tier == 2 and rng.random() < 0.5)
    if has_education:
        level = rng.choices(
            ["bachelor", "master", "phd"],
            weights=[3, 4, 2] if role_key in ("ml", "data_engineer") else [5, 4, 1],
        )[0]
        fld = rng.choice(list(role["field"]))  # type: ignore[index]
        degree = rng.choice(DEGREE_TEMPLATES[level]).format(field=fld)
        school = rng.choice(SCHOOLS)
        grad_year = TODAY_YEAR - years - rng.randint(0, 1)
        education_line = f"{degree}, {school} ({grad_year})"
        education_level: str = level
    else:
        education_level = "unknown"
        education_line = ""

    top = ", ".join(skills[:3])
    summary = (
        f"{display_title} with {years}+ years of experience. "
        f"Core skills include {top} applied across production systems. "
        f"Proven record of delivering results with cross-functional teams."
    )

    jobs = _make_hard_jobs(rng, role, skills, years)
    career_months = _total_months(jobs)

    projects = tuple(
        _fill(p, rng, skills)
        for p in rng.sample(
            list(role["projects"]),  # type: ignore[arg-type]
            min(len(role["projects"]), rng.randint(1, 2)),  # type: ignore[arg-type]
        )
    )
    certs_pool = list(role["certs"])  # type: ignore[index]
    certifications = (
        tuple(rng.sample(certs_pool, min(len(certs_pool), rng.randint(0, 2))))
        if certs_pool else ()
    )

    has_skills_section = (tier != 2)
    # tier-2 ~half have no summary header; others always have one
    has_summary_header = (tier != 2) or (rng.random() >= 0.5)

    if tier == 1:
        exp_header = _rng_choice_from_list(["Career History", "Where I've Worked"], name)
    else:
        exp_header = SECTION_STYLES[rng.choice(list(SECTION_STYLES))]["experience"]

    first, last = name.split(" ", 1)
    return HardResume(
        filename=filename,
        name=name,
        role=role_key,
        title=display_title,
        email=f"{first.lower()}.{last.lower().replace(' ', '')}@example.com",
        phone=f"+1-555-1{100 + idx}",
        location=rng.choice(CITIES),
        summary=summary,
        skills=skills,
        jobs=jobs,
        education_level=education_level,
        education_line=education_line,
        projects=projects,
        certifications=certifications,
        style=rng.choice(list(SECTION_STYLES)),
        order=tuple(rng.choice(SECTION_ORDERS)),
        total_years=career_months // 12,
        tier=tier,
        date_style=date_style,
        present_token=present_token,
        has_skills_section=has_skills_section,
        has_summary_header=has_summary_header,
        experience_header=exp_header,
    )


def _fmt_job_line(job: HardJob, date_style: str, present_token: str) -> str:
    start_str = _fmt_date(job.start_year, job.start_month, date_style, present_token, False)
    end_str = _fmt_date(job.end_year, job.end_month, date_style, present_token, job.is_present)
    return f"{job.title}, {job.company} ({start_str} - {end_str})"


def _hard_plain_text(resume: HardResume) -> str:
    """Render a hard-mode resume as plain text according to its tier."""
    lines: List[str] = []

    if resume.tier == 2:
        # NAME on line 1; rest of contact mashed into line 2
        lines.append(resume.name)
        lines.append(f"{resume.title} | {resume.location} | {resume.email} | {resume.phone}")
    else:
        lines.append(resume.name)
        lines.append(resume.title)
        lines.append(f"{resume.location} | {resume.email} | {resume.phone}")
    lines.append("")

    # Summary block
    if resume.has_summary_header:
        if resume.tier == 1:
            summary_hdr = _rng_choice_from_list(["What I Bring", "My Background"], resume.name)
        else:
            summary_hdr = SECTION_STYLES[resume.style]["summary"]
        lines.append(summary_hdr)
    lines.append(resume.summary)
    lines.append("")

    # Skills (absent for tier 2)
    if resume.has_skills_section:
        if resume.tier == 1:
            lines.append("Toolbox")
        else:
            lines.append(SECTION_STYLES[resume.style]["skills"])
        lines.append(", ".join(resume.skills))
        lines.append("")

    # Experience
    lines.append(resume.experience_header)
    for job in resume.jobs:
        lines.append(_fmt_job_line(job, resume.date_style, resume.present_token))
        lines.extend(f"- {b}" for b in job.bullets)
        lines.append("")

    # Education
    if resume.education_line:
        if resume.tier == 1:
            lines.append("Studies")
        else:
            lines.append(SECTION_STYLES[resume.style]["education"])
        lines.append(resume.education_line)
        lines.append("")

    # Projects
    if resume.projects:
        lines.append(SECTION_STYLES[resume.style]["projects"])
        lines.extend(resume.projects)
        lines.append("")

    # Certifications
    if resume.certifications:
        lines.append(SECTION_STYLES[resume.style]["certifications"])
        for c in resume.certifications:
            lines.append(f"- {c}")
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def _write_hard_txt(resume: HardResume, path: Path) -> None:
    path.write_text(_hard_plain_text(resume), encoding="utf-8")


def _write_hard_pdf(resume: HardResume, path: Path) -> None:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    all_headers: set = set()
    for style in SECTION_STYLES.values():
        all_headers.update(style.values())
    all_headers.update(["Toolbox", "Studies", "Career History", "Where I've Worked",
                         "What I Bring", "My Background"])

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    for i, line in enumerate(_hard_plain_text(resume).splitlines()):
        if not line.strip():
            pdf.ln(3)
            continue
        if i == 0:
            pdf.set_font("Helvetica", style="B", size=15)
        elif line in all_headers:
            pdf.set_font("Helvetica", style="B", size=12)
        else:
            pdf.set_font("Helvetica", size=10)
        pdf.multi_cell(0, 5.5, text=line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.output(str(path))


def _write_hard_docx(resume: HardResume, path: Path) -> None:
    from docx import Document

    doc = Document()
    for i, line in enumerate(_hard_plain_text(resume).splitlines()):
        if i == 0:
            doc.add_heading(line, level=0)
        elif not line.strip():
            continue
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(str(path))


HARD_WRITERS = {".txt": _write_hard_txt, ".pdf": _write_hard_pdf, ".docx": _write_hard_docx}


def _build_hard_label(resume: HardResume) -> Dict[str, object]:
    career_months = _total_months(resume.jobs)
    sy = _skill_years_map(resume.jobs, resume.skills, career_months, resume.date_style)
    return {
        "name": resume.name,
        "role": resume.role,
        "title": resume.title,
        "total_years": resume.total_years,
        "education_level": resume.education_level,
        "skills": sorted(resume.skills),
        "format": Path(resume.filename).suffix.lower(),
        "tier": resume.tier,
        "date_style": resume.date_style,
        "skill_years": {k: v for k, v in sorted(sy.items())},
    }


def _generate_hard_corpus(existing_jd_labels: Dict[str, object]) -> None:
    """Generate 40 hard-mode resumes to resumes_hard/ and labels_hard.json."""
    rng = random.Random(HARD_SEED)
    RESUMES_HARD_DIR.mkdir(parents=True, exist_ok=True)
    DATASET_DIR.mkdir(parents=True, exist_ok=True)

    resume_labels: Dict[str, Dict[str, object]] = {}
    for idx, spec in enumerate(HARD_CANDIDATES):
        resume = _make_hard_resume(idx, spec, rng)
        path = RESUMES_HARD_DIR / resume.filename
        HARD_WRITERS[path.suffix.lower()](resume, path)
        resume_labels[resume.filename] = _build_hard_label(resume)
        print(f"wrote resumes_hard/{resume.filename}")

    labels: Dict[str, object] = {
        "seed": HARD_SEED,
        "today_year": TODAY_YEAR,
        "resumes": resume_labels,
        "job_descriptions": existing_jd_labels,
    }
    labels_path = DATASET_DIR / "labels_hard.json"
    labels_path.write_text(json.dumps(labels, indent=2), encoding="utf-8")
    print(f"\nGenerated {len(resume_labels)} hard resumes")
    print(f"Ground truth written to {labels_path}")


# --- Resume model -------------------------------------------------------------


@dataclass(frozen=True)
class Job:
    title: str
    company: str
    start: int
    end: str  # year as str, or "Present"
    bullets: Tuple[str, ...]


@dataclass(frozen=True)
class Resume:
    filename: str
    name: str
    role: str
    title: str
    email: str
    phone: str
    location: str
    summary: str
    skills: Tuple[str, ...]
    jobs: Tuple[Job, ...]
    education_level: str
    education_line: str
    projects: Tuple[str, ...]
    certifications: Tuple[str, ...]
    style: str
    order: Tuple[str, ...]
    total_years: int


_SENIORITY_WORDS = ("Junior", "Senior", "Staff", "Lead", "Principal")


def _seniority(years: int) -> str:
    if years >= 10:
        return "Staff"
    if years >= 7:
        return "Senior"
    if years <= 2:
        return "Junior"
    return ""


def _with_seniority(title: str, prefix: str) -> str:
    """Prepend a seniority prefix unless the title already carries one."""
    if not prefix or title.split(" ", 1)[0] in _SENIORITY_WORDS:
        return title
    return f"{prefix} {title}"


def _fill(template: str, rng: random.Random, skills: Sequence[str]) -> str:
    s, s2 = rng.sample(list(skills), 2)
    return template.format(s=s, s2=s2, n=rng.randint(1, 9), pct=rng.choice([20, 25, 30, 35, 40, 50, 60, 70]))


def _make_jobs(rng: random.Random, role: Dict[str, object], skills: Sequence[str], years: int) -> Tuple[Job, ...]:
    """Split a `years`-long career into 1-4 contiguous jobs ending at Present."""
    n_jobs = max(1, min(4, (years + 2) // 3))  # ceil(years / 3), capped at 4
    # Correct by construction: every span >= 1 year and the spans sum to `years`.
    spans = [1] * n_jobs
    for _ in range(years - n_jobs):
        spans[rng.randrange(n_jobs)] += 1

    titles = list(role["titles"])  # type: ignore[index]
    bullets = list(role["bullets"])  # type: ignore[index]
    companies = rng.sample(COMPANIES, n_jobs)
    jobs: List[Job] = []
    start = TODAY_YEAR - years
    for i, span in enumerate(spans):
        end_year = start + span
        is_last = i == n_jobs - 1
        end = "Present" if is_last else str(end_year)
        title = rng.choice(titles)
        if is_last:
            title = _with_seniority(title, _seniority(years))
        job_bullets = tuple(_fill(b, rng, skills) for b in rng.sample(bullets, min(2, len(bullets))))
        jobs.append(Job(title=title, company=companies[i], start=start, end=end, bullets=job_bullets))
        start = end_year
    return tuple(reversed(jobs))  # most recent first, like a real resume


def _make_resume(idx: int, spec: Tuple[str, str, str, int], rng: random.Random) -> Resume:
    filename, name, role_key, years = spec
    role = ROLES[role_key]
    core = list(role["core"])  # type: ignore[index]
    extra = list(role["extra"])  # type: ignore[index]
    skills = tuple(core + rng.sample(extra, min(len(extra), rng.randint(2, 4))))

    title = rng.choice(list(role["titles"]))  # type: ignore[index]
    display_title = _with_seniority(title, _seniority(years))

    level = rng.choices(
        ["bachelor", "master", "phd"],
        weights=[3, 4, 2] if role_key in ("ml", "data_engineer") else [5, 4, 1],
    )[0]
    fld = rng.choice(list(role["field"]))  # type: ignore[index]
    degree = rng.choice(DEGREE_TEMPLATES[level]).format(field=fld)
    school = rng.choice(SCHOOLS)
    grad_year = TODAY_YEAR - years - rng.randint(0, 1)
    education_line = f"{degree}, {school} ({grad_year})"

    top = ", ".join(skills[:3])
    summary = (
        f"{display_title} with {years}+ years of experience specialising in {top}. "
        f"Track record of delivering reliable, measurable results in production systems."
    )

    jobs = _make_jobs(rng, role, skills, years)
    projects = tuple(
        _fill(p, rng, skills) for p in rng.sample(list(role["projects"]), min(len(role["projects"]), rng.randint(1, 2)))  # type: ignore[arg-type]
    )
    certs = list(role["certs"])  # type: ignore[index]
    certifications = tuple(rng.sample(certs, min(len(certs), rng.randint(0, 2)))) if certs else ()

    first, last = name.split(" ", 1)
    return Resume(
        filename=filename,
        name=name,
        role=role_key,
        title=display_title,
        email=f"{first.lower()}.{last.lower().replace(' ', '')}@example.com",
        phone=f"+1-555-0{100 + idx}",
        location=rng.choice(CITIES),
        summary=summary,
        skills=skills,
        jobs=jobs,
        education_level=level,
        education_line=education_line,
        projects=projects,
        certifications=certifications,
        style=rng.choice(list(SECTION_STYLES)),
        order=tuple(rng.choice(SECTION_ORDERS)),
        total_years=years,
    )


# --- Renderers ------------------------------------------------------------------


def _section_lines(resume: Resume, section: str) -> List[str]:
    if section == "summary":
        return [resume.summary]
    if section == "skills":
        if resume.style == "title":  # bullet-list variant
            return [f"- {s}" for s in resume.skills]
        return [", ".join(resume.skills)]
    if section == "experience":
        lines: List[str] = []
        for job in resume.jobs:
            lines.append(f"{job.title}, {job.company} ({job.start} - {job.end})")
            lines.extend(f"- {b}" for b in job.bullets)
            lines.append("")
        return lines[:-1] if lines else lines
    if section == "education":
        return [resume.education_line]
    if section == "projects":
        return list(resume.projects)
    if section == "certifications":
        return [f"- {c}" for c in resume.certifications]
    return []


def _plain_text(resume: Resume) -> str:
    headers = SECTION_STYLES[resume.style]
    lines = [resume.name, resume.title, f"{resume.location} | {resume.email} | {resume.phone}", ""]
    for section in resume.order:
        body = _section_lines(resume, section)
        if not body:
            continue
        lines.append(headers[section])
        lines.extend(body)
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def _write_txt(resume: Resume, path: Path) -> None:
    path.write_text(_plain_text(resume), encoding="utf-8")


def _write_pdf(resume: Resume, path: Path) -> None:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    headers = set(SECTION_STYLES[resume.style].values())
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    for i, line in enumerate(_plain_text(resume).splitlines()):
        if not line.strip():
            pdf.ln(3)
            continue
        if i == 0:
            pdf.set_font("Helvetica", style="B", size=15)
        elif line in headers:
            pdf.set_font("Helvetica", style="B", size=12)
        else:
            pdf.set_font("Helvetica", size=10)
        # default new_x leaves the cursor at the right margin, which makes the
        # next multi_cell(0, ...) have zero width
        pdf.multi_cell(0, 5.5, text=line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.output(str(path))


def _write_docx(resume: Resume, path: Path) -> None:
    from docx import Document

    headers = SECTION_STYLES[resume.style]
    doc = Document()
    doc.add_heading(resume.name, level=0)
    doc.add_paragraph(resume.title)
    doc.add_paragraph(f"{resume.location} | {resume.email} | {resume.phone}")
    for section in resume.order:
        body = _section_lines(resume, section)
        if not body:
            continue
        doc.add_heading(headers[section], level=1)
        for line in body:
            if line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line:
                doc.add_paragraph(line)
    doc.save(str(path))


WRITERS = {".txt": _write_txt, ".pdf": _write_pdf, ".docx": _write_docx}


# --- Job descriptions -----------------------------------------------------------

# filename -> (title, body, label) where label feeds dataset/labels.json
JOB_DESCRIPTIONS: Dict[str, Tuple[str, str, Dict[str, object]]] = {
    "senior_ml_engineer.txt": (
        "Senior Machine Learning Engineer",
        """About the role:
Nimbus Pay is hiring a Senior Machine Learning Engineer to own fraud-detection
models end to end, from feature pipelines to monitored production deployments.

Responsibilities:
- Design, train and deploy machine learning models for transaction fraud scoring.
- Build feature pipelines and automated retraining workflows.
- Partner with data engineers to productionise models behind low-latency APIs.

Requirements:
- 5+ years of professional Python experience.
- 3+ years building production machine learning models with PyTorch or TensorFlow.
- Strong grounding in scikit-learn, pandas and SQL.
- Bachelor's degree in Computer Science or a related field.

Nice to have:
- Experience with MLOps tooling, Airflow and AWS.
- Publications or open-source contributions in NLP or Computer Vision.
""",
        {"primary_roles": ["ml"], "adjacent_roles": ["data_engineer"]},
    ),
    "backend_python_engineer.txt": (
        "Backend Engineer (Python)",
        """About the role:
Vertex Logistics needs a Backend Engineer to scale the order-routing platform
that moves freight across three continents.

Responsibilities:
- Build and operate Python microservices with high availability targets.
- Design relational schemas and keep query performance predictable.
- Own services from design through deployment and on-call.

Requirements:
- 4+ years of professional Python experience.
- Production experience with FastAPI or Django.
- Solid PostgreSQL skills including indexing and query tuning.
- Hands-on experience with Docker and REST APIs.

Nice to have:
- Kafka or other event-streaming experience.
- Exposure to Kubernetes and AWS.
""",
        {"primary_roles": ["backend_python"], "adjacent_roles": ["fullstack"]},
    ),
    "frontend_react_engineer.txt": (
        "Frontend Engineer (React)",
        """About the role:
Orchid Media is rebuilding its creator dashboard and needs a Frontend Engineer
who sweats interaction details and performance budgets.

Responsibilities:
- Ship React features with TypeScript across the creator-facing dashboard.
- Evolve our design system and component library.
- Keep Core Web Vitals green on every release.

Requirements:
- 3+ years of professional React experience.
- Strong TypeScript and modern JavaScript fundamentals.
- Deep knowledge of CSS and HTML semantics.

Nice to have:
- Next.js, GraphQL or Tailwind CSS experience.
- Familiarity with Cypress end-to-end testing.
""",
        {"primary_roles": ["frontend"], "adjacent_roles": ["fullstack"]},
    ),
    "devops_platform_engineer.txt": (
        "DevOps / Platform Engineer",
        """About the role:
Clearpath Energy runs workloads on Kubernetes across AWS and GCP and is hiring
a Platform Engineer to make deployments boring.

Responsibilities:
- Own Terraform modules and the GitOps deployment pipeline.
- Operate and upgrade multi-region Kubernetes clusters.
- Build observability with Prometheus and Grafana, and lead incident reviews.

Requirements:
- 4+ years of experience operating Kubernetes in production.
- Strong Terraform and infrastructure-as-code skills.
- Experience with AWS, Docker and CI/CD pipelines.
- Comfortable scripting in Python or Bash on Linux.

Nice to have:
- Certified Kubernetes Administrator (CKA).
- Helm and Ansible experience.
""",
        {"primary_roles": ["devops"], "adjacent_roles": []},
    ),
    "data_platform_engineer.txt": (
        "Data Engineer",
        """About the role:
Northwind Analytics is hiring a Data Engineer to build the pipelines behind our
customer-facing analytics product.

Responsibilities:
- Build and operate batch and streaming pipelines into the warehouse.
- Model analytics tables consumed by dashboards and ML teams.
- Enforce data-quality checks and lineage across sources.

Requirements:
- 3+ years of professional data engineering experience.
- Strong Python and SQL skills.
- Production experience with Spark and Airflow.
- Experience designing and operating production ETL pipelines.

Nice to have:
- Snowflake, dbt or Kafka experience.
- Exposure to machine learning feature pipelines.
""",
        {"primary_roles": ["data_engineer"], "adjacent_roles": ["ml", "backend_python"]},
    ),
    "fullstack_product_engineer.txt": (
        "Full-Stack Product Engineer",
        """About the role:
Mosaic Travel is a small product team shipping weekly. We need a Full-Stack
Engineer who can take features from database schema to polished UI.

Responsibilities:
- Build features across a Python backend and a React + TypeScript frontend.
- Design PostgreSQL schemas and keep APIs fast.
- Collaborate directly with design and customers.

Requirements:
- 3+ years of professional full-stack experience.
- Strong Python and React skills.
- Experience with TypeScript, PostgreSQL and Docker.

Nice to have:
- FastAPI, Next.js or GraphQL experience.
- Experience in early-stage product teams.
""",
        {"primary_roles": ["fullstack"], "adjacent_roles": ["frontend", "backend_python"]},
    ),
}


# --- Main -----------------------------------------------------------------------


def _build_clean_corpus() -> Dict[str, object]:
    """Generate clean corpus; return jd_labels for reuse by hard mode."""
    rng = random.Random(SEED)
    RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    JOBS_DIR.mkdir(parents=True, exist_ok=True)
    DATASET_DIR.mkdir(parents=True, exist_ok=True)

    resume_labels: Dict[str, Dict[str, object]] = {}
    for idx, spec in enumerate(CANDIDATES):
        resume = _make_resume(idx, spec, rng)
        path = RESUMES_DIR / resume.filename
        WRITERS[path.suffix.lower()](resume, path)
        resume_labels[resume.filename] = {
            "name": resume.name,
            "role": resume.role,
            "title": resume.title,
            "total_years": resume.total_years,
            "education_level": resume.education_level,
            "skills": sorted(resume.skills),
            "format": path.suffix.lower(),
        }
        print(f"wrote resumes/{resume.filename}")

    jd_labels: Dict[str, Dict[str, object]] = {}
    for filename, (title, body, label) in JOB_DESCRIPTIONS.items():
        text = f"Job Title: {title}\n\n{body}"
        (JOBS_DIR / filename).write_text(text, encoding="utf-8")
        jd_labels[filename] = {"title": title, **label}
        print(f"wrote job_descriptions/{filename}")

    labels: Dict[str, object] = {
        "seed": SEED,
        "today_year": TODAY_YEAR,
        "resumes": resume_labels,
        "job_descriptions": jd_labels,
    }
    (DATASET_DIR / "labels.json").write_text(json.dumps(labels, indent=2), encoding="utf-8")
    print(f"\nGenerated {len(resume_labels)} resumes, {len(jd_labels)} job descriptions")
    print(f"Ground truth written to {DATASET_DIR / 'labels.json'}")
    return jd_labels


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate labelled resume corpus.")
    parser.add_argument(
        "--hard",
        action="store_true",
        help="Generate ONLY the hard corpus (resumes_hard/ + dataset/labels_hard.json).",
    )
    args = parser.parse_args()

    if args.hard:
        # Load existing jd_labels from labels.json; fall back to inline build.
        labels_path = DATASET_DIR / "labels.json"
        if labels_path.exists():
            existing = json.loads(labels_path.read_text(encoding="utf-8"))
            jd_labels: Dict[str, object] = existing.get("job_descriptions", {})
        else:
            jd_labels = {}
            for filename, (title, _body, label) in JOB_DESCRIPTIONS.items():
                jd_labels[filename] = {"title": title, **label}
        _generate_hard_corpus(jd_labels)
    else:
        _build_clean_corpus()


if __name__ == "__main__":
    main()
