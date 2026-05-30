"""Generate dummy resume files in the ``resumes/`` folder.

Creates a mix of .txt, .pdf and .docx files (8 total) with varied skills so the
read/list/search/summarise tools all have realistic data to work with.

Run from the project root:

    python scripts/make_samples.py
"""

from __future__ import annotations

from pathlib import Path
from typing import Dict, List

PROJECT_ROOT = Path(__file__).resolve().parent.parent
RESUMES_DIR = PROJECT_ROOT / "resumes"

# Each resume: skills/summary vary so keyword searches return different hits.
RESUMES: List[Dict[str, object]] = [
    {
        "filename": "john_doe.pdf",
        "name": "John Doe",
        "title": "Senior Backend Engineer",
        "email": "john.doe@example.com",
        "summary": "Backend engineer with 8 years building Python and Go services.",
        "skills": ["Python", "Go", "PostgreSQL", "Docker", "AWS", "FastAPI"],
        "experience": [
            "Lead Engineer, Acme Cloud (2020-2025): scaled a Python microservice "
            "platform to 2M requests/day.",
            "Backend Developer, DataWorks (2017-2020): built ETL pipelines in Python.",
        ],
    },
    {
        "filename": "jane_smith.docx",
        "name": "Jane Smith",
        "title": "Frontend Engineer",
        "email": "jane.smith@example.com",
        "summary": "Frontend engineer focused on React and TypeScript design systems.",
        "skills": ["JavaScript", "TypeScript", "React", "CSS", "Next.js"],
        "experience": [
            "Senior Frontend Engineer, Pixel Labs (2019-2025): owned the component "
            "library used across 12 products.",
            "UI Developer, Bright Apps (2016-2019): shipped responsive web apps.",
        ],
    },
    {
        "filename": "alex_lee.txt",
        "name": "Alex Lee",
        "title": "Data Scientist",
        "email": "alex.lee@example.com",
        "summary": "Data scientist using Python and PyTorch for NLP and forecasting.",
        "skills": ["Python", "PyTorch", "pandas", "SQL", "scikit-learn"],
        "experience": [
            "Data Scientist, Insight AI (2021-2025): built demand-forecasting models.",
            "ML Intern, Quant Co (2020): prototyped Python notebooks for backtesting.",
        ],
    },
    {
        "filename": "maria_garcia.pdf",
        "name": "Maria Garcia",
        "title": "DevOps Engineer",
        "email": "maria.garcia@example.com",
        "summary": "DevOps engineer automating cloud infrastructure with Terraform.",
        "skills": ["Terraform", "Kubernetes", "AWS", "Python", "CI/CD", "Bash"],
        "experience": [
            "DevOps Lead, Streamline (2019-2025): cut deploy times 70% with GitOps.",
            "SRE, HostPro (2016-2019): managed Kubernetes clusters.",
        ],
    },
    {
        "filename": "sam_patel.txt",
        "name": "Sam Patel",
        "title": "Java Backend Developer",
        "email": "sam.patel@example.com",
        "summary": "Java/Spring Boot developer building enterprise REST APIs.",
        "skills": ["Java", "Spring Boot", "Maven", "MySQL", "Kafka"],
        "experience": [
            "Backend Developer, FinServe (2018-2025): designed Spring Boot services.",
            "Software Engineer, RetailX (2015-2018): maintained Java monolith.",
        ],
    },
    {
        "filename": "lin_wei.docx",
        "name": "Lin Wei",
        "title": "Machine Learning Engineer",
        "email": "lin.wei@example.com",
        "summary": "ML engineer deploying Python models to production at scale.",
        "skills": ["Python", "TensorFlow", "Docker", "Kubernetes", "Airflow"],
        "experience": [
            "ML Engineer, Vision Corp (2020-2025): productionised computer-vision models.",
            "Research Assistant, State University (2018-2020): published 2 papers.",
        ],
    },
    {
        "filename": "olivia_brown.txt",
        "name": "Olivia Brown",
        "title": "Product Designer",
        "email": "olivia.brown@example.com",
        "summary": "Product designer bridging UX research and visual design.",
        "skills": ["Figma", "User Research", "Prototyping", "Design Systems"],
        "experience": [
            "Lead Designer, Craft Studio (2019-2025): owned end-to-end design.",
            "UX Designer, AppWorks (2016-2019): ran usability studies.",
        ],
    },
    {
        "filename": "noah_kim.pdf",
        "name": "Noah Kim",
        "title": "Full-Stack Engineer",
        "email": "noah.kim@example.com",
        "summary": "Full-stack engineer comfortable across Python APIs and React UIs.",
        "skills": ["Python", "Django", "React", "PostgreSQL", "Redis", "Docker"],
        "experience": [
            "Full-Stack Engineer, Loop (2020-2025): built Django + React products.",
            "Junior Developer, WebStart (2018-2020): maintained Python services.",
        ],
    },
]


def _plain_text(resume: Dict[str, object]) -> str:
    """Render a resume as a plain-text document."""
    lines: List[str] = [
        str(resume["name"]),
        str(resume["title"]),
        str(resume["email"]),
        "",
        "SUMMARY",
        str(resume["summary"]),
        "",
        "SKILLS",
        ", ".join(resume["skills"]),  # type: ignore[arg-type]
        "",
        "EXPERIENCE",
    ]
    lines.extend(f"- {item}" for item in resume["experience"])  # type: ignore[union-attr]
    return "\n".join(lines) + "\n"


def _write_txt(resume: Dict[str, object], path: Path) -> None:
    path.write_text(_plain_text(resume), encoding="utf-8")


def _write_pdf(resume: Dict[str, object], path: Path) -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    for line in _plain_text(resume).splitlines():
        # latin-1 is the core-font encoding; sample data is ASCII-safe.
        pdf.cell(0, 8, text=line)
        pdf.ln(8)
    pdf.output(str(path))


def _write_docx(resume: Dict[str, object], path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(str(resume["name"]), level=0)
    doc.add_paragraph(str(resume["title"]))
    doc.add_paragraph(str(resume["email"]))
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(str(resume["summary"]))
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(resume["skills"]))  # type: ignore[arg-type]
    doc.add_heading("Experience", level=1)
    for item in resume["experience"]:  # type: ignore[union-attr]
        doc.add_paragraph(str(item), style="List Bullet")
    doc.save(str(path))


WRITERS = {".txt": _write_txt, ".pdf": _write_pdf, ".docx": _write_docx}


def main() -> None:
    RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    for resume in RESUMES:
        filename = str(resume["filename"])
        ext = Path(filename).suffix.lower()
        path = RESUMES_DIR / filename
        WRITERS[ext](resume, path)
        print(f"wrote {path.relative_to(PROJECT_ROOT)}")
    print(f"\nGenerated {len(RESUMES)} resume files in {RESUMES_DIR}")


if __name__ == "__main__":
    main()
