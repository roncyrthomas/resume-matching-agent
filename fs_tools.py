"""fs_tools.py — Core file-system tools for working with resume files.

Four tools intended for LLM function/tool calling:

    read_file(filepath)                -> dict
    list_files(directory, extension)   -> list
    write_file(filepath, content)      -> dict
    search_in_file(filepath, keyword)  -> dict

Design notes
------------
* Every tool returns a *structured* response and never lets an exception
  escape to the caller. ``read_file``/``write_file``/``search_in_file`` return a
  dict with a ``success`` flag; ``list_files`` returns a list (empty on error)
  to honour the assignment's required signature.
* All file access is confined to a *base directory* (the project root by
  default, or ``$FS_TOOLS_BASE_DIR``). Paths that escape it are rejected, so an
  LLM cannot be coaxed into reading arbitrary system files.
* Text is always decoded as UTF-8 with ``errors="replace"`` — Windows otherwise
  defaults to cp1252 and silently corrupts résumés.
"""

from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple, TypedDict

# --- Constants ---------------------------------------------------------------

SUPPORTED_EXTENSIONS = (".txt", ".pdf", ".docx")
CONTEXT_RADIUS = 40  # characters of context shown on each side of a match


# --- Structured response types ----------------------------------------------


class FileMetadata(TypedDict):
    size: int
    modified: str
    extension: str
    unit: str   # "lines" | "pages" | "paragraphs"
    count: int


class ReadResult(TypedDict, total=False):
    success: bool
    filepath: str
    content: str
    metadata: FileMetadata
    error: str


class WriteResult(TypedDict, total=False):
    success: bool
    filepath: str
    bytes_written: int
    error: str


class SearchMatch(TypedDict, total=False):
    char_start: int
    char_end: int
    snippet: str
    line_number: int


class SearchResult(TypedDict, total=False):
    success: bool
    filepath: str
    keyword: str
    match_count: int
    matches: List[SearchMatch]
    error: str


class FileInfo(TypedDict):
    name: str
    path: str
    size: int
    modified: str
    extension: str


# --- Errors ------------------------------------------------------------------


class FSToolError(Exception):
    """Raised internally for any validation/extraction failure.

    Public tools catch this and translate it into a structured error response.
    """


# --- Path handling -----------------------------------------------------------


def _base_dir() -> Path:
    """Return the sandbox root. Read on every call so tests can override it."""
    return Path(os.environ.get("FS_TOOLS_BASE_DIR", os.getcwd())).resolve()


def _validate_path(user_path: str, *, must_exist: bool, kind: str) -> Path:
    """Resolve *user_path* inside the sandbox and validate it.

    Args:
        user_path: Path supplied by the caller (relative or absolute).
        must_exist: Require the resolved path to exist.
        kind: ``"file"`` or ``"dir"`` — checked only when ``must_exist``.

    Returns:
        The resolved, sandbox-confined :class:`~pathlib.Path`.

    Raises:
        FSToolError: On empty input, sandbox escape, or a failed existence/kind
            check.
    """
    if not user_path or not str(user_path).strip():
        raise FSToolError("path must not be empty")

    base = _base_dir()
    raw = Path(str(user_path))
    candidate = (raw if raw.is_absolute() else base / raw).resolve()

    try:
        candidate.relative_to(base)
    except ValueError as exc:
        raise FSToolError(
            f"path escapes the allowed base directory: {user_path}"
        ) from exc

    if must_exist:
        if not candidate.exists():
            raise FSToolError(f"path does not exist: {user_path}")
        if kind == "file" and not candidate.is_file():
            raise FSToolError(f"not a file: {user_path}")
        if kind == "dir" and not candidate.is_dir():
            raise FSToolError(f"not a directory: {user_path}")

    return candidate


# --- Text extraction ---------------------------------------------------------


def _extract_pdf(path: Path) -> Tuple[str, str, int]:
    """Extract text from a PDF (pdfplumber primary, pypdf fallback)."""
    try:
        import pdfplumber

        pages: List[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        text = "\n".join(pages)
        if text.strip():
            return text, "pages", len(pages)
    except Exception:
        # Fall through to pypdf for damaged files or pdfplumber gaps.
        pass

    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [(pg.extract_text() or "") for pg in reader.pages]
    return "\n".join(pages), "pages", len(reader.pages)


def _extract_docx(path: Path) -> Tuple[str, str, int]:
    """Extract text from a .docx including headers, tables and footers.

    A naive ``paragraph.text`` loop misses text inside tables and
    headers/footers — common in résumé layouts — so we walk all of them.
    """
    from docx import Document

    doc = Document(str(path))
    parts: List[str] = []

    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                parts.append(para.text)

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)

    for section in doc.sections:
        for para in section.footer.paragraphs:
            if para.text.strip():
                parts.append(para.text)

    return "\n".join(parts), "paragraphs", len(doc.paragraphs)


def _extract_text(path: Path) -> Tuple[str, str, int]:
    """Return ``(text, unit, count)`` for a supported file.

    Raises:
        FSToolError: If the extension is unsupported or extraction fails.
    """
    ext = path.suffix.lower()
    try:
        if ext == ".txt":
            text = path.read_text(encoding="utf-8", errors="replace")
            return text, "lines", text.count("\n") + 1 if text else 0
        if ext == ".pdf":
            return _extract_pdf(path)
        if ext == ".docx":
            return _extract_docx(path)
    except FSToolError:
        raise
    except Exception as exc:  # noqa: BLE001 - surface a clean message
        raise FSToolError(f"failed to read {path.name}: {exc}") from exc

    raise FSToolError(
        f"unsupported file type '{ext or '<none>'}'; "
        f"supported: {', '.join(SUPPORTED_EXTENSIONS)}"
    )


def _error(message: str) -> dict:
    return {"success": False, "error": message}


# --- Public tools ------------------------------------------------------------


def read_file(filepath: str) -> ReadResult:
    """Read a resume file (.txt, .pdf or .docx) and extract its text.

    Args:
        filepath: Path to the file, relative to the sandbox root.

    Returns:
        On success: ``{success, filepath, content, metadata}`` where metadata
        holds ``size``, ``modified`` (ISO 8601), ``extension`` and a
        format-dependent ``unit``/``count`` (lines, pages or paragraphs).
        On failure: ``{success: False, error}``.
    """
    try:
        path = _validate_path(filepath, must_exist=True, kind="file")
        text, unit, count = _extract_text(path)
        stat = path.stat()
        return {
            "success": True,
            "filepath": str(path),
            "content": text,
            "metadata": {
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "extension": path.suffix.lower(),
                "unit": unit,
                "count": count,
            },
        }
    except FSToolError as exc:
        return _error(str(exc))
    except Exception as exc:  # noqa: BLE001
        return _error(f"unexpected error reading file: {exc}")


def list_files(directory: str, extension: Optional[str] = None) -> List[FileInfo]:
    """List files in *directory*, optionally filtered by extension.

    Args:
        directory: Directory path relative to the sandbox root.
        extension: Optional filter such as ``".pdf"`` or ``"pdf"`` (the leading
            dot is added automatically).

    Returns:
        A list of ``{name, path, size, modified, extension}`` dicts, sorted by
        name. Returns an empty list if the directory is invalid (graceful
        degradation for the required ``-> list`` signature).
    """
    try:
        path = _validate_path(directory, must_exist=True, kind="dir")
    except FSToolError:
        return []

    if extension and not extension.startswith("."):
        extension = "." + extension
    ext_lower = extension.lower() if extension else None

    results: List[FileInfo] = []
    for child in sorted(path.iterdir(), key=lambda p: p.name.lower()):
        if not child.is_file():
            continue
        if ext_lower and child.suffix.lower() != ext_lower:
            continue
        stat = child.stat()
        results.append(
            {
                "name": child.name,
                "path": str(child),
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "extension": child.suffix.lower(),
            }
        )
    return results


def write_file(filepath: str, content: str) -> WriteResult:
    """Write *content* to *filepath*, creating parent directories as needed.

    Args:
        filepath: Destination path relative to the sandbox root.
        content: Text to write (UTF-8 encoded).

    Returns:
        ``{success, filepath, bytes_written}`` on success, else
        ``{success: False, error}``.
    """
    try:
        if content is None:
            raise FSToolError("content must not be None")
        path = _validate_path(filepath, must_exist=False, kind="file")
        path.parent.mkdir(parents=True, exist_ok=True)
        data = str(content).encode("utf-8")
        path.write_bytes(data)
        return {
            "success": True,
            "filepath": str(path),
            "bytes_written": len(data),
        }
    except FSToolError as exc:
        return _error(str(exc))
    except Exception as exc:  # noqa: BLE001
        return _error(f"unexpected error writing file: {exc}")


def search_in_file(filepath: str, keyword: str) -> SearchResult:
    """Case-insensitively search *filepath* for *keyword* with surrounding text.

    Positions are reported as character offsets into the extracted text, which
    is reliable across all formats (line numbers are meaningless for text
    extracted from PDFs). A ``line_number`` is additionally included for plain
    text and .docx files where it is well defined.

    Args:
        filepath: File to search, relative to the sandbox root.
        keyword: Substring to look for (case-insensitive).

    Returns:
        ``{success, filepath, keyword, match_count, matches}`` where each match
        is ``{char_start, char_end, snippet[, line_number]}``. On failure,
        ``{success: False, error}``.
    """
    try:
        if not keyword or not keyword.strip():
            raise FSToolError("keyword must not be empty")
        path = _validate_path(filepath, must_exist=True, kind="file")
        text, _unit, _count = _extract_text(path)
        track_lines = path.suffix.lower() in (".txt", ".docx")

        matches: List[SearchMatch] = []
        for m in re.finditer(re.escape(keyword), text, re.IGNORECASE):
            start, end = m.start(), m.end()
            ctx_start = max(0, start - CONTEXT_RADIUS)
            ctx_end = min(len(text), end + CONTEXT_RADIUS)
            snippet = text[ctx_start:ctx_end].replace("\n", " ").strip()
            entry: SearchMatch = {
                "char_start": start,
                "char_end": end,
                "snippet": snippet,
            }
            if track_lines:
                entry["line_number"] = text.count("\n", 0, start) + 1
            matches.append(entry)

        return {
            "success": True,
            "filepath": str(path),
            "keyword": keyword,
            "match_count": len(matches),
            "matches": matches,
        }
    except FSToolError as exc:
        return _error(str(exc))
    except Exception as exc:  # noqa: BLE001
        return _error(f"unexpected error searching file: {exc}")
